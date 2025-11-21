# -*- coding: utf-8 -*-
"""
考试服务（业务层）
- 考试会话管理（开始、提交、查看）
- 答案评分
- 成绩报告生成
"""

import os
import json
import time
import uuid
import threading
from typing import Any, Dict, List, Optional, Tuple
from utils import logger

# 考试会话存储
ATTEMPTS_FILE = "./data/exam_attempts.json"
_ATTEMPTS_LOCK = threading.Lock()

def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)

def _ensure_attempts_dir():
    d = os.path.dirname(ATTEMPTS_FILE) or "."
    _ensure_dir(d)

def _now_iso() -> str:
    return time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime())

def _load_attempts() -> Dict[str, Any]:
    """加载所有考试会话"""
    _ensure_attempts_dir()
    if not os.path.exists(ATTEMPTS_FILE):
        return {"attempts": []}
    try:
        with open(ATTEMPTS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"attempts": []}

def _atomic_save_attempts(obj: Dict[str, Any]):
    """原子保存考试会话"""
    _ensure_attempts_dir()
    tmp = ATTEMPTS_FILE + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    os.replace(tmp, ATTEMPTS_FILE)

def _load_bank() -> Dict[str, Any]:
    """加载题库（从 mcq_service 复用）"""
    from services.mcq_service import _load_bank as mcq_load_bank
    return mcq_load_bank()

def _load_paper_questions(paper_id: str) -> List[Dict[str, Any]]:
    """
    根据 paper_id 加载试卷题目
    paper_id 是 ./data/papers 下的文件名（如 "试卷_1234567890.docx"）
    解析 DOCX 文件获取题目
    """
    from services.mcq_service import _read_any_text, parse_mcq_batch_text
    
    paper_path = os.path.join("./data/papers", paper_id)
    if not os.path.exists(paper_path):
        raise FileNotFoundError(f"试卷文件不存在: {paper_id}")
    
    # 读取并解析试卷
    class FileStorage:
        def __init__(self, path):
            self.filename = os.path.basename(path)
            self._path = path
        def read(self):
            with open(self._path, 'rb') as f:
                return f.read()
    
    file_storage = FileStorage(paper_path)
    text, _ = _read_any_text(file_storage)
    items = parse_mcq_batch_text(text)
    
    # 转换为前端需要的格式
    questions = []
    for i, item in enumerate(items, 1):
        # 判断题型（根据答案长度）
        answer = item.get("answer", "")
        qtype = "multi" if len(answer) > 1 else "single"
        
        # 转换选项格式
        options_dict = item.get("options", {})
        options = [{"label": k, "text": v} for k, v in options_dict.items()]
        
        questions.append({
            "qid": f"Q{i:03d}",
            "stem": item.get("stem", ""),
            "qtype": qtype,
            "options": options,
            "answer": answer,  # 标准答案（后端保留，不返回给前端）
            "explain_original": item.get("explain_original", "")
        })
    
    return questions

def papers_list_open() -> List[Dict[str, str]]:
    """
    列出可用的试卷（从 ./data/papers 读取）
    返回格式：[{"paper_id": "xxx.docx", "title": "试卷名", "item_count": 10}]
    """
    from services.mcq_service import bank_list_papers
    result = bank_list_papers()
    
    if not result.get("ok"):
        return []
    
    papers = result.get("papers", [])
    
    # 为每个试卷添加题目数量
    papers_with_count = []
    for paper in papers:
        try:
            questions = _load_paper_questions(paper["paper_id"])
            papers_with_count.append({
                "paper_id": paper["paper_id"],
                "title": paper["title"],
                "item_count": len(questions)
            })
        except Exception as e:
            logger.warning(f"[Exam] 加载试卷失败 {paper['paper_id']}: {e}")
            # 即使加载失败也返回，但题目数为 0
            papers_with_count.append({
                "paper_id": paper["paper_id"],
                "title": paper["title"],
                "item_count": 0
            })
    
    return papers_with_count

def papers_view(paper_id: str) -> Dict[str, Any]:
    """
    查看试卷详情（题目列表，不含答案）
    """
    try:
        questions = _load_paper_questions(paper_id)
        
        # 移除答案和解析（不返回给前端）
        questions_without_answer = []
        for q in questions:
            questions_without_answer.append({
                "qid": q["qid"],
                "stem": q["stem"],
                "qtype": q["qtype"],
                "options": q["options"]
            })
        
        # 提取试卷标题
        title = paper_id
        if title.lower().endswith(".docx"):
            title = title[:-5]
        # 去掉时间戳后缀
        base, sep, tail = title.rpartition("_")
        if base and tail.isdigit():
            title = base
        
        return {
            "ok": True,
            "paper_id": paper_id,
            "title": title,
            "items": questions_without_answer
        }
    except Exception as e:
        logger.error(f"[Exam] 查看试卷失败: {e}", exc_info=True)
        return {"ok": False, "detail": str(e)}

def exam_start(paper_id: str, duration_sec: int, student_id: str = "anonymous") -> Dict[str, Any]:
    """
    开始考试，创建考试会话
    """
    try:
        # 加载试卷题目（含答案，用于后续评分）
        questions = _load_paper_questions(paper_id)
        
        # 创建会话
        attempt_id = uuid.uuid4().hex
        now = _now_iso()
        
        attempt = {
            "attempt_id": attempt_id,
            "paper_id": paper_id,
            "student_id": student_id,
            "duration_sec": duration_sec,
            "start_time": now,
            "end_time": None,
            "questions": questions,  # 保存完整题目（含答案）
            "answers": [],
            "score": None,
            "status": "in_progress"
        }
        
        with _ATTEMPTS_LOCK:
            data = _load_attempts()
            data["attempts"].append(attempt)
            _atomic_save_attempts(data)
        
        return {
            "ok": True,
            "attempt_id": attempt_id,
            "left_sec": duration_sec
        }
    except Exception as e:
        logger.error(f"[Exam] 开始考试失败: {e}", exc_info=True)
        return {"ok": False, "detail": str(e)}

def _calculate_score(question: Dict[str, Any], chosen_labels: List[str]) -> Tuple[float, bool]:
    """
    计算单题得分
    返回：(得分, 是否完全正确)
    """
    std_answer = set(question.get("answer", ""))
    chosen = set(chosen_labels)
    
    if not std_answer:
        return 0.0, False
    
    # 完全正确
    if chosen == std_answer:
        return 1.0, True
    
    # 完全错误或未作答
    if not chosen or chosen.isdisjoint(std_answer):
        return 0.0, False
    
    # 部分正确（多选题）
    if question.get("qtype") == "multi":
        # 选对的数量 / 标准答案数量，但有选错的扣分
        correct_count = len(chosen & std_answer)
        wrong_count = len(chosen - std_answer)
        
        if wrong_count > 0:
            # 有选错的，部分得分
            score = max(0, (correct_count - wrong_count) / len(std_answer))
            return score, False
        else:
            # 只选对了一部分，没有选错的
            score = correct_count / len(std_answer)
            return score, False
    
    # 单选题部分正确视为错误
    return 0.0, False

def exam_submit(attempt_id: str, answers: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    提交答案并评分
    answers: [{"qid": "Q001", "chosen_labels": ["A", "B"]}]
    """
    try:
        with _ATTEMPTS_LOCK:
            data = _load_attempts()
            attempts = data.get("attempts", [])
            
            # 查找会话
            attempt = None
            attempt_idx = None
            for i, a in enumerate(attempts):
                if a.get("attempt_id") == attempt_id:
                    attempt = a
                    attempt_idx = i
                    break
            
            if not attempt:
                return {"ok": False, "detail": "考试会话不存在"}
            
            if attempt.get("status") == "completed":
                return {"ok": False, "detail": "该考试已提交"}
            
            # 构建答案映射
            answer_map = {ans["qid"]: ans["chosen_labels"] for ans in answers}
            
            # 评分
            questions = attempt.get("questions", [])
            items = []
            total_score = 0.0
            
            for q in questions:
                qid = q["qid"]
                chosen = answer_map.get(qid, [])
                score, is_correct = _calculate_score(q, chosen)
                total_score += score
                
                items.append({
                    "qid": qid,
                    "score": score,
                    "is_correct": is_correct
                })
            
            # 更新会话
            attempt["answers"] = answers
            attempt["score"] = total_score
            attempt["items"] = items
            attempt["total_score"] = total_score
            attempt["end_time"] = _now_iso()
            attempt["status"] = "completed"
            
            attempts[attempt_idx] = attempt
            data["attempts"] = attempts
            _atomic_save_attempts(data)
        
        return {
            "ok": True,
            "total_score": total_score,
            "items": items
        }
    except Exception as e:
        logger.error(f"[Exam] 提交答案失败: {e}", exc_info=True)
        return {"ok": False, "detail": str(e)}

def exam_review(attempt_id: str) -> Dict[str, Any]:
    """
    查看答案解析
    """
    try:
        with _ATTEMPTS_LOCK:
            data = _load_attempts()
            attempts = data.get("attempts", [])
            
            # 查找会话
            attempt = None
            for a in attempts:
                if a.get("attempt_id") == attempt_id:
                    attempt = a
                    break
            
            if not attempt:
                return {"ok": False, "detail": "考试会话不存在"}
            
            if attempt.get("status") != "completed":
                return {"ok": False, "detail": "考试尚未提交"}
            
            # 构建答案映射
            answer_map = {ans["qid"]: ans["chosen_labels"] for ans in attempt.get("answers", [])}
            
            # 构建解析数据
            questions = attempt.get("questions", [])
            items = []
            
            for q in questions:
                qid = q["qid"]
                std_answer = q.get("answer", "")
                my_labels = answer_map.get(qid, [])
                
                _, is_correct = _calculate_score(q, my_labels)
                
                items.append({
                    "qid": qid,
                    "stem": q["stem"],
                    "qtype": q["qtype"],
                    "options": q["options"],
                    "correct_labels": list(std_answer),
                    "my_labels": my_labels,
                    "is_correct": is_correct,
                    "analysis": q.get("explain_original", "")
                })
            
            return {
                "ok": True,
                "items": items
            }
    except Exception as e:
        logger.error(f"[Exam] 查看解析失败: {e}", exc_info=True)
        return {"ok": False, "detail": str(e)}

def get_paper_attempts(paper_id: str) -> List[Dict[str, Any]]:
    """
    获取指定试卷的所有已完成考试会话
    """
    with _ATTEMPTS_LOCK:
        data = _load_attempts()
        attempts = data.get("attempts", [])
        
        # 筛选指定试卷且已完成的会话
        paper_attempts = [
            a for a in attempts 
            if a.get("paper_id") == paper_id and a.get("status") == "completed"
        ]
        
        return paper_attempts

def student_export_report_docx(attempt_id: str) -> Dict[str, Any]:
    """
    导出学生成绩报告（DOCX）
    """
    try:
        import docx
    except Exception as e:
        raise RuntimeError("未安装 python-docx，请先安装：pip install python-docx") from e
    
    try:
        with _ATTEMPTS_LOCK:
            data = _load_attempts()
            attempts = data.get("attempts", [])
            
            # 查找会话
            attempt = None
            for a in attempts:
                if a.get("attempt_id") == attempt_id:
                    attempt = a
                    break
            
            if not attempt:
                return {"ok": False, "detail": "考试会话不存在"}
            
            if attempt.get("status") != "completed":
                return {"ok": False, "detail": "考试尚未提交"}
        
        # 生成 DOCX
        doc = docx.Document()
        doc.add_heading("考试成绩报告", level=1)
        
        # 基本信息
        doc.add_paragraph(f"学生ID: {attempt.get('student_id', 'anonymous')}")
        doc.add_paragraph(f"试卷: {attempt.get('paper_id', '')}")
        doc.add_paragraph(f"开始时间: {attempt.get('start_time', '')}")
        doc.add_paragraph(f"结束时间: {attempt.get('end_time', '')}")
        doc.add_paragraph(f"总分: {attempt.get('total_score', 0):.2f}")
        doc.add_paragraph("")
        
        # 构建答案映射
        answer_map = {ans["qid"]: ans["chosen_labels"] for ans in attempt.get("answers", [])}
        
        # 题目详情
        doc.add_heading("答题详情", level=2)
        questions = attempt.get("questions", [])
        
        for i, q in enumerate(questions, 1):
            qid = q["qid"]
            std_answer = q.get("answer", "")
            my_labels = answer_map.get(qid, [])
            _, is_correct = _calculate_score(q, my_labels)
            
            doc.add_paragraph(f"{i}. {q['stem']}")
            
            # 选项
            for opt in q["options"]:
                doc.add_paragraph(f"  {opt['label']}. {opt['text']}")
            
            # 答案
            doc.add_paragraph(f"标准答案: {std_answer}")
            doc.add_paragraph(f"我的答案: {''.join(my_labels) if my_labels else '(未作答)'}")
            doc.add_paragraph(f"判定: {'正确' if is_correct else '错误'}")
            
            # 解析
            analysis = q.get("explain_original", "")
            if analysis:
                doc.add_paragraph(f"解析: {analysis}")
            
            doc.add_paragraph("")
        
        # 保存文件
        _ensure_dir("./data/reports")
        filename = f"成绩报告_{attempt_id[:8]}_{int(time.time())}.docx"
        filepath = os.path.join("./data/reports", filename)
        doc.save(filepath)
        
        return {
            "ok": True,
            "path": filepath,
            "filename": filename,
            "download_url": f"/student/download_report?filename={filename}"
        }
    except Exception as e:
        logger.error(f"[Exam] 导出报告失败: {e}", exc_info=True)
        return {"ok": False, "detail": str(e)}

def export_paper_reports_zip(paper_id: str) -> Tuple[Optional[str], str]:
    """
    导出指定试卷所有考生的成绩报告（ZIP压缩包）
    每个考生一个DOCX文件
    返回：(ZIP文件路径, ZIP文件名)
    """
    try:
        import docx
        import zipfile
    except Exception as e:
        raise RuntimeError("未安装必要的库，请先安装：pip install python-docx") from e
    
    try:
        # 获取该试卷的所有已完成考试
        paper_attempts = get_paper_attempts(paper_id)
        
        if not paper_attempts:
            return None, ""
        
        # 创建临时目录存放各个报告
        _ensure_dir("./data/reports_temp")
        _ensure_dir("./data/reports_zip")
        
        # 生成每个考生的报告
        report_files = []
        for attempt in paper_attempts:
            attempt_id = attempt.get("attempt_id", "")
            student_id = attempt.get("student_id", "anonymous")
            
            # 生成单个报告
            doc = docx.Document()
            doc.add_heading("考试成绩报告", level=1)
            
            # 基本信息
            doc.add_paragraph(f"学生ID: {student_id}")
            doc.add_paragraph(f"试卷: {attempt.get('paper_id', '')}")
            doc.add_paragraph(f"开始时间: {attempt.get('start_time', '')}")
            doc.add_paragraph(f"结束时间: {attempt.get('end_time', '')}")
            doc.add_paragraph(f"总分: {attempt.get('total_score', 0):.2f}")
            doc.add_paragraph("")
            
            # 构建答案映射
            answer_map = {ans["qid"]: ans["chosen_labels"] for ans in attempt.get("answers", [])}
            
            # 题目详情
            doc.add_heading("答题详情", level=2)
            questions = attempt.get("questions", [])
            
            for i, q in enumerate(questions, 1):
                qid = q["qid"]
                std_answer = q.get("answer", "")
                my_labels = answer_map.get(qid, [])
                _, is_correct = _calculate_score(q, my_labels)
                
                doc.add_paragraph(f"{i}. {q['stem']}")
                
                # 选项
                for opt in q["options"]:
                    doc.add_paragraph(f"  {opt['label']}. {opt['text']}")
                
                # 答案
                doc.add_paragraph(f"标准答案: {std_answer}")
                doc.add_paragraph(f"我的答案: {''.join(my_labels) if my_labels else '(未作答)'}")
                doc.add_paragraph(f"判定: {'正确' if is_correct else '错误'}")
                
                # 解析
                analysis = q.get("explain_original", "")
                if analysis:
                    doc.add_paragraph(f"解析: {analysis}")
                
                doc.add_paragraph("")
            
            # 保存单个报告
            filename = f"成绩报告_{student_id}_{attempt_id[:8]}.docx"
            filepath = os.path.join("./data/reports_temp", filename)
            doc.save(filepath)
            report_files.append((filepath, filename))
        
        # 打包成ZIP
        # 提取试卷标题
        title = paper_id
        if title.lower().endswith(".docx"):
            title = title[:-5]
        base, sep, tail = title.rpartition("_")
        if base and tail.isdigit():
            title = base
        
        zip_filename = f"{title}_成绩报告_{int(time.time())}.zip"
        zip_path = os.path.join("./data/reports_zip", zip_filename)
        
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for filepath, filename in report_files:
                zf.write(filepath, arcname=filename)
        
        # 清理临时文件
        for filepath, _ in report_files:
            try:
                os.remove(filepath)
            except Exception:
                pass
        
        return zip_path, zip_filename
        
    except Exception as e:
        logger.error(f"[Exam] 批量导出报告失败: {e}", exc_info=True)
        return None, ""

def export_paper_summary_docx(paper_id: str) -> Tuple[Optional[str], str]:
    """
    导出指定试卷所有考生的成绩汇总表（DOCX）
    包含每名考生的答题情况和总得分
    返回：(DOCX文件路径, DOCX文件名)
    """
    try:
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise RuntimeError("未安装 python-docx，请先安装：pip install python-docx") from e
    
    try:
        # 获取该试卷的所有已完成考试
        paper_attempts = get_paper_attempts(paper_id)
        
        if not paper_attempts:
            return None, ""
        
        # 创建文档
        doc = docx.Document()
        doc.add_heading("考试成绩汇总表", level=1)
        
        # 提取试卷标题
        title = paper_id
        if title.lower().endswith(".docx"):
            title = title[:-5]
        base, sep, tail = title.rpartition("_")
        if base and tail.isdigit():
            title = base
        
        doc.add_paragraph(f"试卷名称: {title}")
        doc.add_paragraph(f"考试人数: {len(paper_attempts)}")
        doc.add_paragraph("")
        
        # 获取题目数量（从第一个attempt中获取）
        if paper_attempts:
            questions = paper_attempts[0].get("questions", [])
            question_count = len(questions)
        else:
            question_count = 0
        
        # 创建表格：学生ID | Q1 | Q2 | ... | Qn | 总分
        table = doc.add_table(rows=1, cols=question_count + 2)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "学生ID"
        for i in range(question_count):
            header_cells[i + 1].text = f"Q{i + 1}"
        header_cells[question_count + 1].text = "总分"
        
        # 设置表头样式
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        
        # 填充每个考生的数据
        for attempt in paper_attempts:
            student_id = attempt.get("student_id", "anonymous")
            total_score = attempt.get("total_score", 0)
            questions = attempt.get("questions", [])
            answer_map = {ans["qid"]: ans["chosen_labels"] for ans in attempt.get("answers", [])}
            
            # 添加新行
            row_cells = table.add_row().cells
            row_cells[0].text = student_id
            
            # 填充每题的答题情况
            for i, q in enumerate(questions):
                qid = q["qid"]
                std_answer = q.get("answer", "")
                my_labels = answer_map.get(qid, [])
                score, is_correct = _calculate_score(q, my_labels)
                
                # 显示：我的答案(正确/错误)
                my_answer_str = ''.join(my_labels) if my_labels else "未答"
                status = "✓" if is_correct else "✗"
                cell_text = f"{my_answer_str} {status}"
                
                row_cells[i + 1].text = cell_text
                row_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 总分
            row_cells[question_count + 1].text = f"{total_score:.2f}"
            row_cells[question_count + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加统计信息
        doc.add_paragraph("")
        doc.add_heading("统计信息", level=2)
        
        # 计算平均分、最高分、最低分
        scores = [a.get("total_score", 0) for a in paper_attempts]
        avg_score = sum(scores) / len(scores) if scores else 0
        max_score = max(scores) if scores else 0
        min_score = min(scores) if scores else 0
        
        doc.add_paragraph(f"平均分: {avg_score:.2f}")
        doc.add_paragraph(f"最高分: {max_score:.2f}")
        doc.add_paragraph(f"最低分: {min_score:.2f}")
        
        # 保存文件
        _ensure_dir("./data/reports")
        filename = f"{title}_成绩汇总_{int(time.time())}.docx"
        filepath = os.path.join("./data/reports", filename)
        doc.save(filepath)
        
        return filepath, filename
        
    except Exception as e:
        logger.error(f"[Exam] 导出成绩汇总失败: {e}", exc_info=True)
        return None, ""
