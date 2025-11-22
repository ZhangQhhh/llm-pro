# -*- coding: utf-8 -*-
"""
选择题服务（业务层，同步版 + 【新增】异步批量与轮询结果回传）
- 保留：同步 /mcq_public/explain 等全部既有能力
- 增强：异步任务在进行中实时累积每题结果（含 strategy / sources / per_option 等）
       get_task_status() 返回 results（累计）以便前端与同步表现对齐
"""

import io
import os
import re
import json
import time
import threading
import uuid
from typing import Any, Dict, List, Optional, Tuple

from flask import current_app
from config import Settings
from utils import logger

# —— 仅 LLM 判定与答案汇总（保持不变） —— #
from utils.mcq.strategy import decide_strategy
from utils.mcq.answer_aggregator import summarize_and_compare

# ===================== 常量与并发保护（保持不变） =====================
BANK_FILE = getattr(Settings, "MCQ_BANK_FILE", "./data/mcq_bank.json")
_BANK_LOCK = threading.Lock()

def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)

def _ensure_bank_dir():
    d = os.path.dirname(BANK_FILE) or "."
    _ensure_dir(d)

def _now_iso() -> str:
    return time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime())

def _load_bank() -> Dict[str, Any]:
    _ensure_bank_dir()
    if not os.path.exists(BANK_FILE):
        return {"seq": 0, "items": []}
    try:
        with open(BANK_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"seq": 0, "items": []}

def _atomic_save_bank(obj: Dict[str, Any]):
    _ensure_bank_dir()
    tmp = BANK_FILE + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
    os.replace(tmp, BANK_FILE)

def _next_id(bank: Dict[str, Any]) -> str:
    bank["seq"] = int(bank.get("seq") or 0) + 1
    return f"Q{bank['seq']:06d}"

def _index_by_id(items: List[Dict[str, Any]]) -> Dict[str, int]:
    return {it["id"]: i for i, it in enumerate(items) if "id" in it}

def _status_from_explain(explain: str) -> str:
    return "draft" if (explain or "").strip() else "none"

# ====== 新增：参考资料持久化（完整保存/读取） ======
import pathlib

SOURCES_DIR = "./data/sources"
def _ensure_sources_dir():
    os.makedirs(SOURCES_DIR, exist_ok=True)

def _sources_path(qid: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_\-\.]", "_", str(qid or ""))
    return os.path.join(SOURCES_DIR, f"{safe}.json")

def save_full_sources(qid: Optional[str], sources: Any):
    if not qid:
        return
    _ensure_sources_dir()
    path = _sources_path(qid)
    tmp  = path + ".tmp"
    try:
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump({"qid": qid, "sources": sources}, f, ensure_ascii=False)
        os.replace(tmp, path)
    except Exception as e:
        logger.warning(f"[MCQ] 保存参考资料失败 qid={qid}: {e}")

def load_full_sources(qid: str) -> Dict[str, Any]:
    _ensure_sources_dir()
    path = _sources_path(qid)
    if not os.path.exists(path):
        return {"ok": True, "qid": qid, "sources": []}
    try:
        with open(path, "r", encoding="utf-8") as f:
            obj = json.load(f)
        # 兼容老文件结构
        src = obj.get("sources", obj)
        return {"ok": True, "qid": qid, "sources": src}
    except Exception as e:
        return {"ok": False, "qid": qid, "msg": f"读取参考资料失败: {e}"}

# 供路由层调用
def bank_get_sources(qid: str) -> Dict[str, Any]:
    if not qid:
        return {"ok": False, "msg": "qid 不能为空"}
    return load_full_sources(qid)

# ===================== 题目解析正则 & 解析器（保持你现有的健壮解析逻辑） =====================
VALID_LABELS = "ABCDEFGH"

OPTION_PAT = re.compile(r"^\s*([A-Ha-h])[\.\)\、：:]\s*(.+?)\s*$")
ANS_LINE_FULL_PAT = re.compile(r"^\s*(?:答案|正确答案|参考答案|【答案】)\s*[:：]?\s*([A-Ha-h]+)\s*(.*)$")
EXPLAIN_LINE_PAT  = re.compile(r"^\s*(?:解析|答案解析|【解析】|知识要点|解答|考点|考点解析|题目解析)\s*[:：]?\s*(.*)$")
ANS_INLINE_PAT    = re.compile(r"(?:^|[；;。.\s])(答案|正确答案|参考答案|正确选项)\s*[:：]?\s*([A-Ha-h]+)(?![A-Za-z])")
SECTION_PREFIX_PAT= re.compile(r"^\s*(?:答案|正确答案|参考答案|解析|答案解析|【答案】|【解析】|知识要点|解答|考点|考点解析|题目解析)\s*[:：]?")
INLINE_EXPLAIN_SPLIT_PAT = re.compile(r"(?:^|\s)(?:解析|答案解析|【解析】|知识要点|解答|考点(?:解析)?|题目解析)\s*[:：]\s*")
QSTART_PAT = re.compile(r"^\s*(?:【第\s*\d+\s*题.*?】|第?\s*\d+\s*题|[0-9]+\s*[\.、)])\s*", re.M)
STEM_PREFIX_PAT = re.compile(r"""^\s*(?:【第\s*\d+\s*题.*?】|第?\s*\d+\s*题|[（(]?\d+[)）]|[0-9]+\s*[\.、：:])\s*""", re.X)
BRACKET_HEADER_PAT = re.compile(r"^\s*【第\s*(\d+)\s*题\s*[-—]\s*(单选|多选)】\s*(.*)$")

def strip_q_prefix(s: str) -> str:
    return STEM_PREFIX_PAT.sub("", (s or "").strip())

def _read_docx_text(b: bytes) -> str:
    try:
        import docx
    except Exception as e:
        raise RuntimeError("未安装 python-docx，请先安装：pip install python-docx") from e
    f = io.BytesIO(b)
    doc = docx.Document(f)
    return "\n".join(p.text for p in doc.paragraphs)

def _read_any_text(file_storage) -> Tuple[str, str]:
    filename = (file_storage.filename or "unknown").lower()
    data = file_storage.read()
    if filename.endswith(".txt"):
        text = data.decode("utf-8", errors="ignore")
    elif filename.endswith(".docx"):
        text = _read_docx_text(data)
    else:
        try:
            text = data.decode("utf-8")
        except Exception:
            raise RuntimeError("暂不支持的文件类型，请上传 .txt 或 .docx")
    return text, filename.rsplit(".", 1)[-1] if "." in filename else ""

def split_mcq_blocks(all_text: str) -> List[str]:
    text = all_text.replace("\r\n", "\n").replace("\r", "\n")
    starts = [m.start() for m in QSTART_PAT.finditer(text)]
    if starts:
        starts.append(len(text))
        blocks = [text[starts[i]:starts[i+1]].strip() for i in range(len(starts)-1)]
    else:
        parts = re.split(r"\n\s*\n\s*\n+", text)
        blocks = [p.strip() for p in parts if p.strip()]
    filtered = [b for b in blocks if re.search(r"(?m)^\s*[A-Ha-h][\.\)\、：:]\s*\S+", b)]
    return filtered or (blocks or [all_text])

def _clean_answer(ans: str) -> str:
    letters = [c.upper() for c in ans if c.upper() in VALID_LABELS]
    return "".join(sorted(dict.fromkeys(letters)))

def _split_option_inline_explain(opt_text: str) -> Tuple[str, str]:
    m = INLINE_EXPLAIN_SPLIT_PAT.search(opt_text or "")
    if not m:
        return (opt_text or "").strip(), ""
    pure = (opt_text[:m.start()] or "").rstrip()
    tail = (opt_text[m.end():] or "").strip()
    return pure, tail

def parse_single_mcq_text(text: str) -> Dict[str, Any]:
    lines = [ln.rstrip() for ln in text.split("\n")]
    while lines and not lines[0].strip(): lines.pop(0)
    while lines and not lines[-1].strip(): lines.pop()

    stem_lines: List[str] = []
    explain_lines: List[str] = []
    options: Dict[str, str] = {}
    answer = ""
    in_options = False
    in_explain = False
    cur_label = ""

    # 识别【第x题-单选|多选】头
    if lines:
        m_hdr = BRACKET_HEADER_PAT.match(lines[0])
        if m_hdr:
            tail_stem = (m_hdr.group(3) or "").strip()
            if tail_stem:
                stem_lines.append(tail_stem)
            lines = lines[1:]

    def _flush_cur_label():
        nonlocal cur_label
        if cur_label and cur_label in options:
            options[cur_label] = options[cur_label].strip()
        cur_label = ""

    for raw in lines:
        line = raw.rstrip()

        m_exp = EXPLAIN_LINE_PAT.match(line)
        if m_exp:
            in_explain = True
            in_options = False
            _flush_cur_label()
            tail = (m_exp.group(1) or "").strip()
            if tail:
                explain_lines.append(tail)
            continue

        if in_explain:
            m_ans_inside = ANS_LINE_FULL_PAT.match(line)
            if m_ans_inside:
                if not answer:
                    answer = _clean_answer(m_ans_inside.group(1))
                tail = (m_ans_inside.group(2) or "").strip()
                if tail:
                    explain_lines.append(tail)
                continue
            if not answer:
                mi = ANS_INLINE_PAT.search(line)
                if mi:
                    answer = _clean_answer(mi.group(2))
            explain_lines.append(line)
            continue

        m_ans = ANS_LINE_FULL_PAT.match(line)
        if m_ans:
            in_options = False
            _flush_cur_label()
            answer = _clean_answer(m_ans.group(1) or answer)
            tail = (m_ans.group(2) or "").strip()
            if tail:
                in_explain = True
                explain_lines.append(tail)
            continue

        m_opt = OPTION_PAT.match(line)
        if m_opt:
            in_options = True
            in_explain = False
            _flush_cur_label()
            lab = m_opt.group(1).upper()
            val = (m_opt.group(2) or "").strip()
            pure, tail_exp = _split_option_inline_explain(val)
            options[lab] = pure
            cur_label = lab
            if tail_exp:
                in_explain = True
                explain_lines.append(tail_exp)
                _flush_cur_label()
                in_options = False
            continue

        if in_options and cur_label:
            if SECTION_PREFIX_PAT.match(line):
                _flush_cur_label()
                in_options = False
            else:
                if line.strip():
                    options[cur_label] = (options[cur_label] + " " + line.strip()).strip()
            continue

        stem_lines.append(line)

    _flush_cur_label()

    # D 选项尾随“解析”尾巴清理
    if not explain_lines and options:
        last_lab = None
        for k in VALID_LABELS:
            if k in options: last_lab = k
        if last_lab:
            pure, tail = _split_option_inline_explain(options[last_lab])
            if tail:
                options[last_lab] = pure
                explain_lines.append(tail)

    # 答案兜底：从解析或题干行中提取
    if not answer:
        mi = ANS_INLINE_PAT.search("\n".join(explain_lines))
        if mi:
            answer = _clean_answer(mi.group(2))
        else:
            mi2 = ANS_INLINE_PAT.search("\n".join(stem_lines))
            if mi2:
                answer = _clean_answer(mi2.group(2))

    stem_raw = "\n".join([s for s in stem_lines if s.strip()]).strip()
    stem = strip_q_prefix(stem_raw or (lines[0].strip() if lines else ""))

    clean_opts: Dict[str, str] = {}
    for k in VALID_LABELS:
        if k in options and (options[k] or "").strip():
            clean_opts[k] = options[k].strip()
        else:
            if clean_opts:
                break

    explain_original = "\n".join([s for s in explain_lines if str(s).strip()]).strip()

    return {
        "stem": stem,
        "options": clean_opts,
        "answer": (answer or "").upper(),
        "explain_original": explain_original,
    }

def parse_mcq_batch_text(all_text: str) -> List[Dict[str, Any]]:
    blocks = split_mcq_blocks(all_text)
    items = []
    for i, b in enumerate(blocks, 1):
        parsed = parse_single_mcq_text(b)
        items.append({
            "qid": i,
            "stem": parsed["stem"],
            "options": parsed["options"],
            "answer": parsed["answer"],
            "explain_original": parsed["explain_original"],
            "raw_block": b
        })
    return items

# ===================== 组装问句（与你前端保持完全一致） =====================
def compose_qa_question(stem: str, options: Dict[str, str]) -> str:
    lines = [f"问题：{stem.strip()}", "候选答案："]
    for k in "ABCDEFGH":
        if k in options and options[k]:
            lines.append(f"{k}. {options[k]}")
    return "\n".join(lines).strip()

def compose_qa_question_single(stem: str, label: str, text: str) -> str:
    return f"问题：{stem.strip()}\n候选答案：\n{label}. {text.strip()}"

# ===================== 同步解析（保持不变） =====================
def _std_answer_by_qid_from_bank(qid: Optional[str]) -> str:
    if not qid:
        return ""
    with _BANK_LOCK:
        bank = _load_bank()
        idx = _index_by_id(bank["items"])
        i = idx.get(str(qid))
        if i is None:
            return ""
        return (bank["items"][i].get("answer") or "").upper()

def _collect_stream(process_iter, thinking: bool):
    explain_parts, think_parts, sources = [], [], []
    for item in process_iter:
        if isinstance(item, tuple) and len(item) == 2:
            prefix, payload = item
            if prefix == "THINK" and isinstance(payload, str):
                if thinking: think_parts.append(payload)
            elif prefix == "CONTENT" and isinstance(payload, str):
                explain_parts.append(payload)
            elif prefix == "SOURCE":
                try: sources.append(json.loads(payload))
                except Exception: pass
        else:
            s = str(item)
            if s.startswith("SOURCE:"):
                try: sources.append(json.loads(s[7:]))
                except Exception: pass
            elif s.startswith("THINK:"):
                if thinking: think_parts.append(s[6:])
            elif s.startswith("CONTENT:"):
                explain_parts.append(s[8:])
            else:
                explain_parts.append(s)
    return "".join(explain_parts).strip(), (think_parts if thinking else []), sources

def _process_one_question(
    qid: Optional[str],
    stem: str,
    options: Dict[str, str],
    *,
    model_id: Optional[str],
    thinking: bool,
    rerank_top_n: int,
    classify_on: bool,
    use_insert_block: bool,
    insert_block_llm_id: Optional[str],
) -> Dict[str, Any]:

    llm_service = getattr(current_app, "llm_service", None)
    knowledge_handler = getattr(current_app, "knowledge_handler", None)
    if llm_service is None or knowledge_handler is None:
        raise RuntimeError("服务未初始化：缺少 llm_service / knowledge_handler")

    requested_model_id = model_id or Settings.DEFAULT_LLM_ID
    strategy_model_id = getattr(Settings, "STRATEGY_LLM_ID", None) or requested_model_id

    if classify_on:
        try:
            strategy_client = llm_service.get_client(strategy_model_id)
            strategy, meta = decide_strategy(stem, options, llm_client=strategy_client, temperature=0.0)
            logger.info(f"[MCQ][Strategy] qid={qid} => {strategy}")
        except Exception as e:
            logger.warning(f"[MCQ][StrategyFail] qid={qid}: {e}")
            strategy = "SIMPLE_LOOKUP"
    else:
        strategy = "SIMPLE_LOOKUP"

    llm_client = llm_service.get_client(requested_model_id)

    if strategy == "SIMPLE_LOOKUP":
        composed = compose_qa_question(stem, options)
        explain, thinks, sources = _collect_stream(
            knowledge_handler.process(
                composed, thinking, (int(rerank_top_n) or 20), llm_client, client_ip="sync",
                use_insert_block=use_insert_block, insert_block_llm_id=insert_block_llm_id
            ),
            thinking,
        )
        std_answer = _std_answer_by_qid_from_bank(qid)
        agg = summarize_and_compare(
            stem, options, std_answer, llm_client,
            strategy="SIMPLE_LOOKUP", simple_explain=explain, temperature=0.0
        )
        explain = (explain + ("\n" + agg["summary_block"] if agg["summary_block"] else "")).strip()
        save_full_sources(qid, sources)
        return {
            "qid": qid, "ok": True, "strategy": strategy, "per_option": [],
            "explain": explain, "sources": sources,
            "final_answer": agg["final_answer"],
            "std_answer": agg["std_answer"],
            "answer_mismatch": agg["answer_mismatch"],
        }

    # 复杂：逐项验证
    per_option, sources_grouped = [], []
    for lab in "ABCDEFGH":
        if lab not in options or not (options[lab] or "").strip():
            continue
        composed = compose_qa_question_single(stem, lab, options[lab])
        ex, th, opt_src = _collect_stream(
            knowledge_handler.process(
                composed, thinking, (int(rerank_top_n) or 20), llm_client, client_ip="sync",
                use_insert_block=use_insert_block, insert_block_llm_id=insert_block_llm_id
            ),
            thinking,
        )
        per_option.append({"label": lab, "composed": composed, "explain": ex, "sources": opt_src})
        sources_grouped.append({"label": lab, "sources": opt_src})

    summary_lines = ["【复杂验证（逐选项核查·汇总）】", f"问题：{stem}", "", "分项解析："]
    for it in per_option:
        ex = (it.get("explain") or "").strip() or "（无解析）"
        summary_lines.append(f"{it['label']}. {ex}")
    summary_lines += ["", "说明：本步骤仅对分项结果做汇总，不进行额外推断与判定。"]
    explain_total = "\n".join(summary_lines).strip()

    std_answer = _std_answer_by_qid_from_bank(qid)
    agg = summarize_and_compare(
        stem, options, std_answer, llm_client,
        strategy="COMPLEX_VALIDATION", per_option=per_option, temperature=0.0
    )
    explain_total = (explain_total + ("\n" + agg["summary_block"] if agg["summary_block"] else "")).strip()
    # ---> 新增：保存完整参考资料（分组选项）
    save_full_sources(qid, sources_grouped)
    return {
        "qid": qid, "ok": True, "strategy": "COMPLEX_VALIDATION",
        "per_option": per_option, "sources": sources_grouped,
        "explain": explain_total,
        "final_answer": agg["final_answer"],
        "std_answer": agg["std_answer"],
        "answer_mismatch": agg["answer_mismatch"],
    }

def mcq_explain_sync(data: Dict[str, Any]) -> Dict[str, Any]:
    items = data.get("items")
    enable_thinking = bool(data.get("thinking", False))
    requested_model_id = data.get("model_id") or Settings.DEFAULT_LLM_ID
    use_insert_block = bool(data.get("use_insert_block", False))
    insert_block_llm_id = data.get("insert_block_llm_id")
    rerank_top_n = int(data.get("rerank_top_n") or getattr(Settings, "RERANK_TOP_N", 20))

    classify_global = bool(getattr(Settings, "MCQ_ENABLE_CLASSIFY_STRATEGY", True))
    enable_classify = data.get("enable_classify_strategy")
    if enable_classify is None:
        enable_classify = classify_global

    task_list: List[Dict[str, Any]] = []
    if items and isinstance(items, list):
        for it in items:
            stem = strip_q_prefix((it.get("stem") or "").strip())
            if not stem: continue
            task_list.append({"qid": it.get("qid") or it.get("id"), "stem": stem, "options": it.get("options") or {}})
    else:
        stem = strip_q_prefix((data.get("stem") or "").strip())
        if not stem: return {"ok": False, "msg": "stem 不能为空"}
        task_list.append({"qid": data.get("qid") or data.get("id"), "stem": stem, "options": data.get("options") or {}})

    results = []
    for task in task_list:
        try:
            res = _process_one_question(
                task.get("qid"), task["stem"], task["options"],
                model_id=requested_model_id,
                thinking=enable_thinking,
                rerank_top_n=rerank_top_n,
                classify_on=bool(enable_classify),
                use_insert_block=use_insert_block,
                insert_block_llm_id=insert_block_llm_id,
            )
            results.append({"ok": True, **res})
        except Exception as e:
            logger.error(f"[MCQ] 同步解析异常 qid={task.get('qid')}: {e}", exc_info=True)
            results.append({"ok": False, "qid": task.get("qid"), "msg": str(e)})

    return {"ok": True, "count": len(results), "results": results}

# ===================== 上传 & 题库 CRUD（保持不变） =====================
def mcq_upload_parse(file_storage) -> Dict[str, Any]:
    text, ext = _read_any_text(file_storage)
    items = parse_mcq_batch_text(text)
    return {"ok": True, "filename": file_storage.filename or "", "ext": ext, "count": len(items), "items": items}

def bank_list() -> Dict[str, Any]:
    with _BANK_LOCK:
        bank = _load_bank()
        # 默认不返回已删除的题目
        items = [it for it in bank.get("items", []) if not it.get("deleted")]
        return {"ok": True, "count": len(items), "items": items}

def bank_bulk_upsert(data: Dict[str, Any]) -> Dict[str, Any]:
    items = data.get("items") or []
    if not isinstance(items, list) or not items:
        return {"ok": False, "msg": "items 不能为空"}
    now = _now_iso()
    out: List[Dict[str, Any]] = []

    with _BANK_LOCK:
        bank = _load_bank()
        idx = _index_by_id(bank["items"])

        for it in items:
            options = it.get("options") or {}
            if not isinstance(options, dict):
                options = {}
            answer = (it.get("answer") or "").upper()
            explain = it.get("explain") or ""
            status = it.get("status") or _status_from_explain(explain)

            rec = {
                "id": it.get("id") or "",
                "stem": it.get("stem") or "",
                "options": options,
                "answer": answer,
                "explain": explain,
                "status": status,
                "created_at": it.get("created_at") or now,
                "updated_at": now
            }

            if not rec["id"]:
                rec["id"] = _next_id(bank)
                bank["items"].append(rec)
                idx[rec["id"]] = len(bank["items"]) - 1
            else:
                if rec["id"] in idx:
                    bank["items"][idx[rec["id"]]] = rec
                else:
                    bank["items"].append(rec)
                    idx[rec["id"]] = len(bank["items"]) - 1
            out.append(rec)

        _atomic_save_bank(bank)
        return {"ok": True, "count": len(out), "items": out}

def bank_bulk_update(data: Dict[str, Any]) -> Dict[str, Any]:
    items = data.get("items") or []
    if not isinstance(items, list) or not items:
        return {"ok": False, "msg": "items 不能为空"}

    now = _now_iso()
    updated: List[Dict[str, Any]] = []
    with _BANK_LOCK:
        bank = _load_bank()
        idx = _index_by_id(bank["items"])
        for it in items:
            _id = it.get("id")
            if not _id or _id not in idx:
                continue
            j = idx[_id]
            rec = dict(bank["items"][j])
            if "stem" in it and it["stem"] is not None:
                rec["stem"] = it["stem"]
            if "options" in it and isinstance(it["options"], dict):
                new_opts = dict(rec.get("options") or {})
                for k, v in (it["options"] or {}).items():
                    new_opts[k] = v
                rec["options"] = new_opts
            if "answer" in it:
                rec["answer"] = (it.get("answer") or "").upper()
            if "explain" in it:
                rec["explain"] = it.get("explain") or ""
            if "status" in it and it["status"] is not None:
                rec["status"] = it["status"]
            else:
                rec["status"] = _status_from_explain(rec.get("explain") or "")
            rec["updated_at"] = now
            bank["items"][j] = rec
            updated.append(rec)

        _atomic_save_bank(bank)
        return {"ok": True, "count": len(updated), "items": updated}

def bank_bulk_reject(data: Dict[str, Any]) -> Dict[str, Any]:
    ids = data.get("ids") or []
    if not ids and isinstance(data.get("items"), list):
        ids = [str(it.get("id")) for it in (data.get("items") or []) if it.get("id")]
    ids = [str(x).strip() for x in ids if str(x).strip()]
    if not ids:
        return {"ok": False, "msg": "ids 不能为空"}

    reason = (data.get("reason") or "").strip()
    now = _now_iso()
    updated = []
    with _BANK_LOCK:
        bank = _load_bank()
        idx = _index_by_id(bank.get("items", []))
        for _id in ids:
            j = idx.get(_id)
            if j is None:
                continue
            rec = dict(bank["items"][j])
            rec["status"] = "rejected"
            if reason:
                rec["reject_reason"] = reason
            rec["updated_at"] = now
            bank["items"][j] = rec
            updated.append(rec)
        _atomic_save_bank(bank)

    return {"ok": True, "count": len(updated), "items": updated}

def bank_delete_questions(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    软删除题目（移到回收站）
    data: {"ids": ["qid1", "qid2", ...], "user": "username", "permanent": false}
    """
    ids = data.get("ids") or []
    if not ids and isinstance(data.get("items"), list):
        ids = [str(it.get("id")) for it in (data.get("items") or []) if it.get("id")]
    ids = [str(x).strip() for x in ids if str(x).strip()]
    if not ids:
        return {"ok": False, "msg": "ids 不能为空"}

    user = data.get("user", "unknown")
    permanent = data.get("permanent", False)
    now = _now_iso()
    deleted = []
    
    with _BANK_LOCK:
        bank = _load_bank()
        items = bank.get("items", [])
        idx = _index_by_id(items)
        
        if permanent:
            # 永久删除：直接从列表中移除
            for _id in ids:
                j = idx.get(_id)
                if j is not None:
                    deleted.append(items[j])
            
            bank["items"] = [it for it in items if str(it.get("id")) not in ids]
            _log_deletion(ids, user, "permanent", deleted)
        else:
            # 软删除：标记为已删除
            for _id in ids:
                j = idx.get(_id)
                if j is None:
                    continue
                rec = dict(bank["items"][j])
                rec["deleted"] = True
                rec["deleted_at"] = now
                rec["deleted_by"] = user
                bank["items"][j] = rec
                deleted.append(rec)
            
            _log_deletion(ids, user, "soft", deleted)
        
        _atomic_save_bank(bank)

    return {"ok": True, "count": len(deleted), "deleted": deleted}

def bank_list_deleted() -> Dict[str, Any]:
    """列出回收站中的题目"""
    with _BANK_LOCK:
        bank = _load_bank()
        items = bank.get("items", [])
        deleted_items = [it for it in items if it.get("deleted") == True]
        return {"ok": True, "count": len(deleted_items), "items": deleted_items}

def bank_restore_questions(data: Dict[str, Any]) -> Dict[str, Any]:
    """从回收站恢复题目"""
    ids = data.get("ids") or []
    if not ids and isinstance(data.get("items"), list):
        ids = [str(it.get("id")) for it in (data.get("items") or []) if it.get("id")]
    ids = [str(x).strip() for x in ids if str(x).strip()]
    if not ids:
        return {"ok": False, "msg": "ids 不能为空"}

    user = data.get("user", "unknown")
    now = _now_iso()
    restored = []
    
    with _BANK_LOCK:
        bank = _load_bank()
        items = bank.get("items", [])
        idx = _index_by_id(items)
        
        for _id in ids:
            j = idx.get(_id)
            if j is None:
                continue
            rec = dict(bank["items"][j])
            if rec.get("deleted") == True:
                rec["deleted"] = False
                rec["restored_at"] = now
                rec["restored_by"] = user
                # 移除删除相关字段
                rec.pop("deleted_at", None)
                rec.pop("deleted_by", None)
                bank["items"][j] = rec
                restored.append(rec)
        
        _atomic_save_bank(bank)
        _log_restoration(ids, user, restored)

    return {"ok": True, "count": len(restored), "restored": restored}

def bank_clear_deleted(data: Dict[str, Any]) -> Dict[str, Any]:
    """清空回收站（永久删除所有已删除的题目）"""
    user = data.get("user", "unknown")
    days = data.get("days", 30)  # 默认删除30天前的
    now_ts = time.time()
    cutoff_ts = now_ts - (days * 24 * 3600)
    
    with _BANK_LOCK:
        bank = _load_bank()
        items = bank.get("items", [])
        
        to_remove = []
        for it in items:
            if it.get("deleted") != True:
                continue
            deleted_at = it.get("deleted_at", "")
            if not deleted_at:
                continue
            # 解析时间
            try:
                dt = time.strptime(deleted_at, "%Y-%m-%dT%H:%M:%S")
                ts = time.mktime(dt)
                if ts < cutoff_ts:
                    to_remove.append(str(it.get("id")))
            except:
                continue
        
        if to_remove:
            removed_items = [it for it in items if str(it.get("id")) in to_remove]
            bank["items"] = [it for it in items if str(it.get("id")) not in to_remove]
            _atomic_save_bank(bank)
            _log_deletion(to_remove, user, "auto_clear", removed_items)
            return {"ok": True, "count": len(to_remove), "removed": removed_items}
        
    return {"ok": True, "count": 0, "removed": []}

def _log_deletion(ids: List[str], user: str, action: str, items: List[Dict]) -> None:
    """记录删除日志"""
    try:
        log_dir = "./data/logs"
        _ensure_dir(log_dir)
        log_file = os.path.join(log_dir, "deletion_log.json")
        
        log_entry = {
            "timestamp": _now_iso(),
            "user": user,
            "action": action,  # soft, permanent, auto_clear
            "ids": ids,
            "count": len(ids),
            "items": [{"id": it.get("id"), "stem": it.get("stem", "")[:50]} for it in items]
        }
        
        logs = []
        if os.path.exists(log_file):
            try:
                with open(log_file, "r", encoding="utf-8") as f:
                    logs = json.load(f)
            except:
                logs = []
        
        logs.append(log_entry)
        
        # 只保留最近1000条日志
        if len(logs) > 1000:
            logs = logs[-1000:]
        
        with open(log_file, "w", encoding="utf-8") as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"[MCQ] 记录删除日志失败: {e}")

def _log_restoration(ids: List[str], user: str, items: List[Dict]) -> None:
    """记录恢复日志"""
    try:
        log_dir = "./data/logs"
        _ensure_dir(log_dir)
        log_file = os.path.join(log_dir, "deletion_log.json")
        
        log_entry = {
            "timestamp": _now_iso(),
            "user": user,
            "action": "restore",
            "ids": ids,
            "count": len(ids),
            "items": [{"id": it.get("id"), "stem": it.get("stem", "")[:50]} for it in items]
        }
        
        logs = []
        if os.path.exists(log_file):
            try:
                with open(log_file, "r", encoding="utf-8") as f:
                    logs = json.load(f)
            except:
                logs = []
        
        logs.append(log_entry)
        
        if len(logs) > 1000:
            logs = logs[-1000:]
        
        with open(log_file, "w", encoding="utf-8") as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"[MCQ] 记录恢复日志失败: {e}")

def bank_get_deletion_logs(limit: int = 100) -> Dict[str, Any]:
    """获取删除日志"""
    try:
        log_file = "./data/logs/deletion_log.json"
        if not os.path.exists(log_file):
            return {"ok": True, "logs": []}
        
        with open(log_file, "r", encoding="utf-8") as f:
            logs = json.load(f)
        
        # 返回最近的N条
        logs = logs[-limit:] if len(logs) > limit else logs
        logs.reverse()  # 最新的在前
        
        return {"ok": True, "logs": logs, "count": len(logs)}
    except Exception as e:
        return {"ok": False, "msg": f"读取日志失败: {e}"}

# 导入导出/试卷（保持不变）
def bank_export_docx() -> Tuple[str, str]:
    try:
        import docx
    except Exception as e:
        raise RuntimeError("未安装 python-docx，请先安装：pip install python-docx") from e

    with _BANK_LOCK:
        bank = _load_bank()
        items = bank.get("items", [])

    doc = docx.Document()
    doc.add_heading("题库导出", level=1)
    for i, it in enumerate(items, 1):
        stem = it.get("stem") or ""
        doc.add_paragraph(f"{i}. {stem}")
        opts = it.get("options") or {}
        for k in "ABCDEFGH":
            if k in opts and opts[k]:
                doc.add_paragraph(f"{k}. {opts[k]}")
        ans = (it.get("answer") or "").upper()
        if ans:
            doc.add_paragraph(f"答案：{ans}")
        exp = (it.get("explain") or "").strip()
        if exp:
            doc.add_paragraph(f"解析：{exp}")
        doc.add_paragraph("")
    _ensure_dir("./data/export")
    path = f"./data/export/bank_export_{int(time.time())}.docx"
    doc.save(path)
    return path, os.path.basename(path)

def bank_import_docx(file_storage) -> Dict[str, Any]:
    text, _ = _read_any_text(file_storage)
    items = parse_mcq_batch_text(text)
    now = _now_iso()
    with _BANK_LOCK:
        bank = {"seq": 0, "items": []}
        for it in items:
            rec = {
                "id": _next_id(bank),
                "stem": it["stem"],
                "options": it["options"],
                "answer": (it.get("answer") or "").upper(),
                "explain": it.get("explain_original") or "",
                "status": _status_from_explain(it.get("explain_original") or ""),
                "created_at": now,
                "updated_at": now,
            }
            bank["items"].append(rec)
        _atomic_save_bank(bank)
    return {"ok": True, "count": len(items)}

def bank_generate_paper(name: str) -> Tuple[Optional[str], str]:
    try:
        import docx
    except Exception as e:
        raise RuntimeError("未安装 python-docx，请先安装：pip install python-docx") from e

    with _BANK_LOCK:
        bank = _load_bank()
        approved = [it for it in bank.get("items", []) if it.get("status") == "approved"]

    if not approved:
        return None, f"{name}.docx"

    doc = docx.Document()
    doc.add_heading(name, level=1)
    for i, it in enumerate(approved, 1):
        stem = it.get("stem") or ""
        doc.add_paragraph(f"{i}. {stem}")
        opts = it.get("options") or {}
        for k in "ABCDEFGH":
            if k in opts and opts[k]:
                doc.add_paragraph(f"{k}. {opts[k]}")
        
        # 添加答案
        ans = (it.get("answer") or "").upper()
        if ans:
            doc.add_paragraph(f"答案：{ans}")
        
        # 添加解析
        exp = (it.get("explain") or "").strip()
        if exp:
            doc.add_paragraph(f"解析：{exp}")
        
        doc.add_paragraph("")
    _ensure_dir("./data/papers")
    path = f"./data/papers/{name}_{int(time.time())}.docx"
    doc.save(path)
    return path, os.path.basename(path)

def bank_list_papers() -> Dict[str, Any]:
    """列出 ./data/papers 下已生成的试卷，适配前端的 Paper 接口。"""
    base_dir = "./data/papers"
    try:
        if not os.path.isdir(base_dir):
            return {"ok": True, "papers": []}
        papers: List[Dict[str, str]] = []
        for fname in sorted(os.listdir(base_dir)):
            if not fname.lower().endswith(".docx"):
                continue
            full_path = os.path.join(base_dir, fname)
            if not os.path.isfile(full_path):
                continue
            # 默认用文件名作为 paper_id，便于下载
            paper_id = fname
            title = fname
            # 去掉扩展名
            if title.lower().endswith(".docx"):
                title = title[:-5]
            # name_时间戳 形式时，仅保留 name 作为标题
            base, sep, tail = title.rpartition("_")
            if base and tail.isdigit():
                title = base
            papers.append({"paper_id": paper_id, "title": title})
        # 按文件名倒序（一般包含时间戳），最近的在前面
        papers.sort(key=lambda p: p["paper_id"], reverse=True)
        return {"ok": True, "papers": papers}
    except Exception as e:
        logger.error(f"[MCQ] 列出试卷失败: {e}", exc_info=True)
        return {"ok": False, "msg": str(e)}


def _resolve_paper_path(paper_id: str) -> Optional[str]:
    """根据 paper_id 解析试卷文件路径。
    目前 paper_id 直接使用文件名；如果传入不带 .docx，也做一次容错匹配。
    """
    base_dir = "./data/papers"
    if not paper_id:
        return None
    if not os.path.isdir(base_dir):
        return None

    # 1) 直接按文件名匹配
    direct = os.path.join(base_dir, paper_id)
    if os.path.isfile(direct):
        return direct

    # 2) 不带后缀时，自动补 .docx
    if not paper_id.lower().endswith(".docx"):
        candidate = os.path.join(base_dir, paper_id + ".docx")
        if os.path.isfile(candidate):
            return candidate

    return None


def bank_get_paper_docx(paper_id: str) -> Tuple[Optional[str], str]:
    """获取单个试卷 DOCX 的路径与下载文件名。"""
    path = _resolve_paper_path(paper_id)
    if not path:
        return None, ""
    filename = os.path.basename(path)
    return path, filename


def bank_get_paper_zip(paper_id: str) -> Tuple[Optional[str], str]:
    """将单个试卷打包成 ZIP，返回 ZIP 路径与文件名。"""
    import zipfile

    docx_path, filename = bank_get_paper_docx(paper_id)
    if not docx_path:
        return None, ""

    _ensure_dir("./data/papers_zip")
    zip_name = os.path.splitext(filename)[0] + ".zip"
    zip_path = os.path.join("./data/papers_zip", zip_name)

    try:
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(docx_path, arcname=filename)
    except Exception as e:
        logger.error(f"[MCQ] 打包试卷 ZIP 失败: {e}", exc_info=True)
        return None, ""
    return zip_path, zip_name

# ===================== 【增强】异步任务实现（返回累计结果） =====================
_TASKS: Dict[str, Dict[str, Any]] = {}
_TASKS_LOCK = threading.Lock()

def _task_set(task_id: str, **kv):
    with _TASKS_LOCK:
        cur = _TASKS.get(task_id, {})
        cur.update(kv)
        cur["updated_at"] = _now_iso()
        _TASKS[task_id] = cur
        return dict(cur)

def _task_get(task_id: str) -> Optional[Dict[str, Any]]:
    with _TASKS_LOCK:
        cur = _TASKS.get(task_id)
        if not cur:
            return None
        # 返回深拷贝，避免外部修改内部状态
        out = dict(cur)
        if "results" in out:
            out["results"] = list(out["results"])
        return out

def _task_append_result(task_id: str, res: Dict[str, Any]):
    """将单题结果追加到任务结果列表（只保留前端需要的关键字段）"""
    safe = {
        "qid": res.get("qid"),
        "ok": bool(res.get("ok", True)),
        "strategy": res.get("strategy"),
        "sources": res.get("sources") or [],
        "per_option": res.get("per_option") or [],
        "explain": res.get("explain") or "",
        "final_answer": res.get("final_answer"),
        "std_answer": res.get("std_answer"),
        "answer_mismatch": res.get("answer_mismatch"),
    }
    with _TASKS_LOCK:
        cur = _TASKS.get(task_id, {})
        lst = list(cur.get("results") or [])
        lst.append(safe)
        cur["results"] = lst
        cur["updated_at"] = _now_iso()
        _TASKS[task_id] = cur

def create_async_explain_task(data: Dict[str, Any]) -> Dict[str, Any]:
    """创建一个异步任务（选取 无解析/已驳回/异常 的题目进行批量解析），并在任务状态中累计返回每题结果。"""
    enable_thinking = bool(data.get("thinking", False))
    requested_model_id = data.get("model_id") or Settings.DEFAULT_LLM_ID
    use_insert_block = bool(data.get("use_insert_block", False))
    insert_block_llm_id = data.get("insert_block_llm_id")
    rerank_top_n = int(data.get("rerank_top_n") or getattr(Settings, "RERANK_TOP_N", 20))

    # 策略分类开关（延续全局/请求开关）
    classify_global = bool(getattr(Settings, "MCQ_ENABLE_CLASSIFY_STRATEGY", True))
    enable_classify = data.get("enable_classify_strategy")
    if enable_classify is None:
        enable_classify = classify_global

    # 选择待处理清单
    with _BANK_LOCK:
        bank = _load_bank()
        todo = [it for it in bank.get("items", []) if (it.get("status") or "none") in ("none", "rejected", "abnormal")]

    task_id = uuid.uuid4().hex
    _task_set(task_id, id=task_id, status="queued", done=0, total=len(todo), created_at=_now_iso(), results=[])

    # 后台线程：逐题处理并写回题库，同时累计结果
    app = current_app._get_current_object()

    def _worker():
        with app.app_context():
            try:
                _task_set(task_id, status="running")
                done = 0
                for it in todo:
                    qid = it.get("id")
                    try:
                        res = _process_one_question(
                            qid, it.get("stem") or "", it.get("options") or {},
                            model_id=requested_model_id,
                            thinking=enable_thinking,
                            rerank_top_n=rerank_top_n,
                            classify_on=bool(enable_classify),
                            use_insert_block=use_insert_block,
                            insert_block_llm_id=insert_block_llm_id,
                        )
                        # 累计结果（供前端轮询展示参考资料/分项）
                        _task_append_result(task_id, {"ok": True, **res})
                        save_full_sources(qid, res.get("sources"))
                        # 写回解析与状态
                        explain = (res.get("explain") or "").strip()
                        mismatch = bool(res.get("answer_mismatch"))
                        new_status = "abnormal" if mismatch else "draft"
                        bank_bulk_update({"items": [{"id": qid, "explain": explain, "status": new_status}]})
                    except Exception as e:
                        logger.error(f"[MCQ][Async] 题目处理失败 qid={qid}: {e}", exc_info=True)
                        _task_append_result(task_id, {"ok": False, "qid": qid, "msg": str(e)})
                    finally:
                        done += 1
                        _task_set(task_id, done=done)
                _task_set(task_id, status="done")
            except Exception as e:
                logger.error(f"[MCQ][Async] 任务异常: {e}", exc_info=True)
                _task_set(task_id, status="failed")

    th = threading.Thread(target=_worker, name=f"mcq_async_{task_id}", daemon=True)
    th.start()

    return {"ok": True, "task_id": task_id, "total": len(todo)}

def get_task_status(task_id: str) -> Dict[str, Any]:
    t = _task_get(task_id)
    if not t:
        return {"ok": False, "msg": "任务不存在"}
    # 返回累计结果（前端将与同步模式一致地渲染参考资料）
    return {
        "ok": True,
        "task_id": task_id,
        "status": t.get("status"),
        "done": int(t.get("done") or 0),
        "total": int(t.get("total") or 0),
        "results": list(t.get("results") or []),
        "updated_at": t.get("updated_at"),
    }


