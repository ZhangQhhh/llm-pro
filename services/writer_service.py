# -*- coding: utf-8 -*-
"""
WriterService（统一会话集合版）
- 会话知识库：统一写入单集合（Settings.WRITER_SESSION_COLLECTION）；仍支持 TTL / 覆盖上传 / 精删
- 额外知识库：持久化集合不变（Settings.WRITER_KB_COLLECTION）
- 检索：会话与KB独立topk + 来源配额 → RRF 合并 → 最终裁剪
- 切分：沿用 core.document_processor.split_documents_writer（写作友好）
- 兼容：attach_files / attach_kb_files 接受可选参数
- 新增：retrieve 增参 kb_selected（list[str]，仅使用被选中的 KB 文档；元素按 basename 比对）
"""

import os
import time
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple, Any

from qdrant_client import QdrantClient
from qdrant_client.http import models as qmodels
from core.custom_qdrant_store import FixedQdrantVectorStore as QdrantVectorStore

from llama_index.core import Document, StorageContext, VectorStoreIndex, QueryBundle
from llama_index.core.schema import TextNode, NodeWithScore

from core.retriever import RetrieverFactory
from core.document_processor import split_documents_writer
from utils.logger import logger
from config import Settings as AppSettings


# ------- 文本读取器 -------
def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def _read_md(path: str) -> str:
    return _read_txt(path)

def _read_pdf(path: str) -> str:
    text = []
    try:
        from pypdf import PdfReader
    except Exception:
        from PyPDF2 import PdfReader
    try:
        reader = PdfReader(path)
        for p in reader.pages:
            try:
                text.append(p.extract_text() or "")
            except Exception:
                pass
    except Exception as e:
        logger.warning(f"读取 PDF 失败：{path} | {e}")
    return "\n".join(text)

def _read_docx(path: str) -> str:
    try:
        import docx
    except ImportError as e:
        raise RuntimeError("缺少依赖 python-docx，无法解析 .docx 文件") from e
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs)

def _read_doc(path: str) -> str:
    try:
        import textract
        return textract.process(path).decode("utf-8", errors="ignore")
    except Exception as e:
        raise RuntimeError("解析 .doc 需要把文件另存为 .docx") from e

READERS = {".txt": _read_txt, ".md": _read_md, ".pdf": _read_pdf, ".docx": _read_docx, ".doc": _read_doc}
ALLOWED_EXTS = set(READERS.keys())


# ------- 数据结构 -------
@dataclass
class WriterSessionIndex:
    docs: List[Document] = field(default_factory=list)
    nodes: List[TextNode] = field(default_factory=list)     # 供 BM25 用
    index: Optional[VectorStoreIndex] = None                # 会话内“就地”索引（仅当前会话 nodes）
    collection: Optional[str] = None                        # 统一会话集合名
    expires_at: Optional[float] = None                      # 过期时间戳（秒）

@dataclass
class WriterKBIndex:
    docs: List[Document] = field(default_factory=list)
    nodes: List[TextNode] = field(default_factory=list)     # 供 BM25 用
    index: Optional[VectorStoreIndex] = None
    collection: Optional[str] = None                        # Qdrant 持久化集合名


# ------- 服务主体 -------
class WriterService:
    def __init__(self):
        self._sessions: Dict[str, WriterSessionIndex] = {}
        self._kb: WriterKBIndex = WriterKBIndex()

        self.qdrant_client = QdrantClient(
            host=getattr(AppSettings, "QDRANT_HOST", "localhost"),
            port=int(getattr(AppSettings, "QDRANT_PORT", 6333))
        )

        self.use_qdrant: bool = bool(getattr(AppSettings, "WRITER_USE_QDRANT", True))
        default_kb = f"{getattr(AppSettings, 'QDRANT_COLLECTION', 'kb')}_writer"
        self.kb_collection: str = str(getattr(AppSettings, "WRITER_KB_COLLECTION", default_kb))

        default_sess = f"{getattr(AppSettings, 'QDRANT_COLLECTION', 'kb')}_writer_session"
        self.session_collection: str = str(getattr(AppSettings, "WRITER_SESSION_COLLECTION", default_sess))

        self.sess_persist: bool = bool(getattr(AppSettings, "WRITER_SESSION_PERSIST", True))
        self.sess_ttl_hours: int = int(getattr(AppSettings, "WRITER_SESSION_TTL_HOURS", 24))

        logger.info(
            f"[WriterDBG] qdrant: use={self.use_qdrant}, "
            f"kb_collection={self.kb_collection}, session_collection={self.session_collection}, "
            f"sess_persist={self.sess_persist}, ttl_hours={self.sess_ttl_hours}"
        )

    @staticmethod
    def allowed_ext(path: str) -> bool:
        return os.path.splitext(path)[1].lower() in ALLOWED_EXTS

    @staticmethod
    def _ensure_embed():
        from llama_index.core import Settings as _S
        if getattr(_S, "embed_model", None) is None:
            raise RuntimeError("全局 Embedding 未初始化，请确认 Settings.embed_model 已注入")

    def _ensure_session(self, session_id: str) -> WriterSessionIndex:
        if session_id not in self._sessions:
            self._sessions[session_id] = WriterSessionIndex(collection=self.session_collection)
        if not self._sessions[session_id].collection:
            self._sessions[session_id].collection = self.session_collection
        return self._sessions[session_id]

    def _vector_store(self, collection: str) -> QdrantVectorStore:
        return QdrantVectorStore(client=self.qdrant_client, collection_name=collection)

    @staticmethod
    def _name_from_meta(meta: Dict) -> str:
        return meta.get("filename") or meta.get("file_name") or os.path.basename(meta.get("path") or meta.get("file_path") or "")

    # —— 更安全的删除：可选 session_id 过滤 —— #
    def _qdrant_delete_by_filename(self, collection: str, filename: str, session_id: Optional[str] = None) -> int:
        try:
            should = [
                qmodels.FieldCondition(key="metadata.filename", match=qmodels.MatchValue(value=filename)),
                qmodels.FieldCondition(key="metadata.file_name", match=qmodels.MatchValue(value=filename)),
                qmodels.FieldCondition(key="filename", match=qmodels.MatchValue(value=filename)),
                qmodels.FieldCondition(key="file_name", match=qmodels.MatchValue(value=filename)),
            ]
            must = []
            if session_id:
                must.append(qmodels.FieldCondition(key="metadata.session_id", match=qmodels.MatchValue(value=str(session_id))))
            flt = qmodels.Filter(must=must, should=should) if must else qmodels.Filter(should=should)
            res = self.qdrant_client.delete(collection_name=collection, points_selector=flt)
            logger.info(f"[WriterDBG] qdrant_delete_by_filename: collection={collection}, filename={filename}, session_id={session_id}, status={getattr(res, 'status', 'ok')}")
            return 0
        except Exception as e:
            logger.warning(f"[WriterDBG] qdrant_delete_by_filename failed: {collection} | {filename} | sid={session_id} | {e}")
            return 0

    def _load_nodes_from_qdrant(self, collection: str):
        try:
            logger.info(f"[WriterDBG] qdrant_load_nodes: collection={collection}")
            next_page = None
            out = []
            while True:
                points, next_page = self.qdrant_client.scroll(
                    collection_name=collection,
                    scroll_filter=None,
                    limit=512,
                    with_payload=True,
                    with_vectors=False,
                    offset=next_page
                )
                for p in points or []:
                    payload = p.payload or {}
                    meta = payload.get("metadata") or {}
                    # —— 直接当纯文本取 —— #
                    text = payload.get("_node_content") or payload.get("text") or ""
                    if not text:
                        continue
                    # 补齐常见元数据键
                    for k in ("filename", "file_name", "path", "file_path", "scope", "session_id", "chunk_id"):
                        if k in payload and k not in meta:
                            meta[k] = payload[k]
                    out.append(TextNode(text=text, metadata=meta))
                if next_page is None:
                    break
            logger.info(f"[WriterDBG] qdrant_load_nodes: fetched={len(out)}")
            return out
        except Exception as e:
            logger.warning(f"[WriterDBG] qdrant_load_nodes failed: {collection} | {e}")
            return []

    # ====== 会话文件（上传/删除/列出） ======
    def attach_files(self, session_id: str, saved_file_paths: List[str],
                     target_collection: Optional[str] = None,
                     ingest_metadata: Optional[dict] = None,
                     **kwargs: Any) -> Tuple[int, List[str]]:
        self._ensure_embed()
        sess = self._ensure_session(session_id)
        if target_collection:
            sess.collection = target_collection

        added, names = 0, []
        for p in saved_file_paths:
            ext = os.path.splitext(p)[1].lower()
            if ext not in ALLOWED_EXTS:
                logger.warning(f"不支持的文件类型：{p}")
                continue
            try:
                filename = os.path.basename(p)
                # 覆盖同名（内存）
                if sess.docs:
                    sess.docs = [d for d in sess.docs if (d.metadata.get("filename") or d.metadata.get("file_name")) != filename]
                if sess.nodes:
                    sess.nodes = [n for n in sess.nodes if self._name_from_meta(n.metadata) != filename]
                # 覆盖同名（Qdrant，单集合 + session_id）
                if self.use_qdrant and self.sess_persist and sess.collection:
                    self._qdrant_delete_by_filename(sess.collection, filename, session_id=session_id)

                text = READERS[ext](p)
                if not text.strip():
                    logger.warning(f"文件为空或无法提取文本：{p}")
                    continue

                meta = {
                    "filename": filename,
                    "file_name": filename,
                    "path": p,
                    "file_path": p,
                    "session_id": session_id
                }
                if ingest_metadata:
                    meta.update({k: v for k, v in ingest_metadata.items() if k not in meta})
                doc = Document(text=text, metadata=meta)
                sess.docs.append(doc); added += 1; names.append(filename)
            except Exception as e:
                logger.warning(f"解析失败：{p} | {e}")
                raise

        if added > 0:
            sess.index = None
        return added, names

    def detach_file(self, session_id: str, filename: str) -> bool:
        sess = self._ensure_session(session_id)
        removed = False; keep: List[Document] = []
        for d in sess.docs:
            if (d.metadata.get("filename") == filename) or (d.metadata.get("file_name") == filename):
                removed = True
                try:
                    path = d.metadata.get("path") or d.metadata.get("file_path")
                    if path and os.path.exists(path): os.remove(path)
                except Exception:
                    pass
            else:
                keep.append(d)
        if removed:
            sess.docs = keep
            if sess.nodes:
                sess.nodes = [n for n in sess.nodes if self._name_from_meta(n.metadata) != filename]
            if self.use_qdrant and self.sess_persist and sess.collection:
                self._qdrant_delete_by_filename(sess.collection, filename, session_id=session_id)
            sess.index = None
        return removed

    def list_files(self, session_id: str) -> List[str]:
        return [d.metadata.get("filename") or d.metadata.get("file_name", "") for d in self._ensure_session(session_id).docs]

    def get_doc_text_by_filename(self, session_id: str, filename: str) -> Optional[str]:
        for d in self._ensure_session(session_id).docs:
            if (d.metadata.get("filename") == filename) or (d.metadata.get("file_name") == filename):
                return d.text
        return None

    # ====== 额外知识库（持久化 + 列出 + 覆盖） ======
    def attach_kb_files(self, saved_file_paths: List[str],
                        target_collection: Optional[str] = None,
                        ingest_metadata: Optional[dict] = None,
                        **kwargs: Any) -> Tuple[int, List[str]]:
        self._ensure_embed()
        kb = self._kb
        if not kb.collection:
            kb.collection = target_collection or self.kb_collection

        added, names = 0, []
        for p in saved_file_paths:
            ext = os.path.splitext(p)[1].lower()
            if ext not in ALLOWED_EXTS:
                logger.warning(f"[KB] 不支持的文件类型：{p}")
                continue
            try:
                filename = os.path.basename(p)
                # 覆盖（内存）
                if kb.docs:
                    kb.docs = [d for d in kb.docs if (d.metadata.get("filename") or d.metadata.get("file_name")) != filename]
                if kb.nodes:
                    kb.nodes = [n for n in kb.nodes if self._name_from_meta(n.metadata) != filename]
                # 覆盖（Qdrant）
                if self.use_qdrant:
                    self._qdrant_delete_by_filename(kb.collection, filename)

                text = READERS[ext](p)
                if not text.strip():
                    logger.warning(f"[KB] 文件空或无法提取文本：{p}")
                    continue

                meta = {
                    "filename": filename,
                    "file_name": filename,
                    "path": p,
                    "file_path": p,
                    "scope": "writer_kb"
                }
                if ingest_metadata:
                    meta.update({k: v for k, v in ingest_metadata.items() if k not in meta})
                doc = Document(text=text, metadata=meta)
                kb.docs.append(doc); added += 1; names.append(filename)
            except Exception as e:
                logger.warning(f"[KB] 解析失败：{p} | {e}")
                raise

        if added > 0:
            kb.index = None  # 触发重建（切分 + 入库）
        return added, names

    def detach_kb_file(self, filename: str) -> bool:
        kb = self._kb
        removed = False; keep: List[Document] = []
        for d in kb.docs:
            if (d.metadata.get("filename") == filename) or (d.metadata.get("file_name") == filename):
                removed = True
                try:
                    path = d.metadata.get("path") or d.metadata.get("file_path")
                    if path and os.path.exists(path): os.remove(path)
                except Exception:
                    pass
            else:
                keep.append(d)
        if removed:
            kb.docs = keep
            if kb.nodes:
                kb.nodes = [n for n in kb.nodes if self._name_from_meta(n.metadata) != filename]
            if self.use_qdrant and kb.collection:
                self._qdrant_delete_by_filename(kb.collection, filename)
            kb.index = None
        return removed

    def list_kb_files(self) -> List[str]:
        kb = self._kb
        names = set()
        for d in kb.docs:
            names.add(d.metadata.get("filename") or d.metadata.get("file_name") or "")
        for n in kb.nodes or []:
            nm = self._name_from_meta(n.metadata)
            if nm: names.add(nm)
        if self.use_qdrant:
            collection = kb.collection or self.kb_collection
            try:
                next_page = None
                while True:
                    resp = self.qdrant_client.scroll(
                        collection_name=collection,
                        scroll_filter=None,
                        limit=512,
                        with_payload=True,
                        with_vectors=False,
                        offset=next_page
                    )
                    points, next_page = resp[0], resp[1]
                    for p in points or []:
                        payload = p.payload or {}
                        meta = payload.get("metadata") or {}
                        nm = meta.get("filename") or meta.get("file_name") or payload.get("filename") or payload.get("file_name")
                        if not nm:
                            path_like = meta.get("path") or meta.get("file_path") or payload.get("path") or payload.get("file_path")
                            if path_like:
                                nm = os.path.basename(path_like)
                        if nm:
                            names.add(nm)
                    if next_page is None:
                        break
            except Exception as e:
                logger.warning(f"[WriterDBG] list_kb_files scan failed: {collection} | {e}")

        out = sorted([n for n in names if n])
        logger.info(f"[WriterDBG] list_kb_files: count={len(out)}")
        return out

    # ====== 建索引：会话 ======
    def _build_session_index_if_needed(self, session_id: str):
        sess = self._ensure_session(session_id)

        if self.sess_ttl_hours > 0 and sess.expires_at:
            if time.time() > sess.expires_at:
                self._sessions[session_id] = WriterSessionIndex(collection=self.session_collection)
                sess = self._sessions[session_id]

        if sess.index is not None:
            return
        if not sess.docs:
            return

        split_docs: List[Document] = split_documents_writer(
            sess.docs, max_chars=1200, overlap=200, prefer_headings=True, merge_short=True
        )

        nodes: List[TextNode] = []
        for i, d in enumerate(split_docs):
            meta = dict(d.metadata or {})
            meta.setdefault("filename", meta.get("file_name") or os.path.basename(meta.get("path", "")) or "")
            meta.setdefault("file_name", meta.get("filename", ""))
            meta.setdefault("path", meta.get("file_path", ""))
            meta.setdefault("file_path", meta.get("path", ""))
            meta.setdefault("session_id", session_id)
            meta["chunk_id"] = meta.get("chunk_id", i)
            if self.sess_ttl_hours > 0:
                expire_at = time.time() + self.sess_ttl_hours * 3600
                meta["expires_at"] = expire_at
                sess.expires_at = expire_at
            nodes.append(TextNode(text=d.text, metadata=meta))

        sess.nodes = nodes

        if self.use_qdrant and self.sess_persist and sess.collection:
            vs = self._vector_store(sess.collection)
            ctx = StorageContext.from_defaults(vector_store=vs)
            _ = VectorStoreIndex(nodes, storage_context=ctx)  # upsert

        ctx_local = StorageContext.from_defaults()
        sess.index = VectorStoreIndex(nodes, storage_context=ctx_local)
        logger.info(f"[WriterDBG] build_session_index: local_index nodes={len(nodes)} | persisted_to={sess.collection if (self.use_qdrant and self.sess_persist) else 'memory'}")

    # ====== 建索引：额外知识库 ======
    def _build_kb_index_if_needed(self):
        kb = self._kb

        if kb.index is not None:
            return

        if not kb.collection:
            kb.collection = self.kb_collection

        if kb.docs:
            split_docs: List[Document] = split_documents_writer(
                kb.docs, max_chars=1200, overlap=200, prefer_headings=True, merge_short=True
            )
            nodes: List[TextNode] = []
            for i, d in enumerate(split_docs):
                meta = dict(d.metadata or {})
                meta.setdefault("filename", meta.get("file_name") or os.path.basename(meta.get("path", "")) or "")
                meta.setdefault("file_name", meta.get("filename", ""))
                meta.setdefault("path", meta.get("file_path", ""))
                meta.setdefault("file_path", meta.get("path", ""))
                meta.setdefault("scope", "writer_kb")
                meta["chunk_id"] = meta.get("chunk_id", i)
                nodes.append(TextNode(text=d.text, metadata=meta))

            kb.nodes = kb.nodes + nodes

            if self.use_qdrant:
                vs = self._vector_store(kb.collection)
                ctx = StorageContext.from_defaults(vector_store=vs)
                _ = VectorStoreIndex(nodes, storage_context=ctx)  # upsert
                kb.index = VectorStoreIndex.from_vector_store(vs)
                logger.info(f"[WriterDBG] build_kb_index: persisted to qdrant collection={kb.collection}, add_nodes={len(nodes)}, total_nodes={len(kb.nodes)}")
            else:
                ctx = StorageContext.from_defaults()
                kb.index = VectorStoreIndex(kb.nodes, storage_context=ctx)
                logger.info(f"[WriterDBG] build_kb_index: in-memory, nodes={len(kb.nodes)}")

            kb.docs = []
            return

        if self.use_qdrant:
            vs = self._vector_store(kb.collection)
            kb.index = VectorStoreIndex.from_vector_store(vs)
            if not kb.nodes:
                kb.nodes = self._load_nodes_from_qdrant(kb.collection)
            logger.info(f"[WriterDBG] build_kb_index: restored from qdrant collection={kb.collection}, nodes={len(kb.nodes)}")

    # ====== 安全裁剪 ======
    @staticmethod
    def _safe_k(requested: int, total: int, default_min: int = 1) -> int:
        if total <= 0:
            return 0
        try:
            k = int(requested)
        except Exception:
            k = default_min
        if k < 1:
            k = default_min
        if k > total:
            k = total
        return k

    @staticmethod
    def _peek_nodes(tag: str, items: List[NodeWithScore], n: int = 3):
        try:
            samples = []
            for i, nw in enumerate(items[:n], start=1):
                meta = getattr(nw, "node", nw).metadata if hasattr(nw, "node") else {}
                fn = (meta.get("filename") or meta.get("file_name") or (meta.get("path") or "").split("/")[-1] or "")
                samples.append(f"{i}:{fn}#chunk{meta.get('chunk_id','?')}")
            logger.info(f"[WriterDBG] {tag} peek: " + (", ".join(samples) if samples else "(empty)"))
        except Exception:
            pass

    # ====== 检索主流程（支持 kb_selected 筛选） ======
    def retrieve(self, session_id: str, query: str, use_kb: bool = True,
                 kb_selected: Optional[List[str]] = None) -> List[NodeWithScore]:
        # 构建/恢复索引
        self._build_session_index_if_needed(session_id)
        self._build_kb_index_if_needed()

        sess = self._sessions.get(session_id)
        kb = self._kb

        # —— 选中文件集合（basename 比对） —— #
        selected_kb_set = set()
        if use_kb and kb_selected:
            for x in kb_selected:
                if x:
                    selected_kb_set.add(os.path.basename(str(x)))

        # 各源独立 top_k（回退链条与原一致）
        req_sim_sess = int(getattr(
            AppSettings, "WRITER_SESSION_SIM_TOP_K",
            getattr(AppSettings, "WRITER_SIMILARITY_TOP_K",
                    getattr(AppSettings, "RETRIEVAL_TOP_K", 30))
        ))
        req_bm25_sess = int(getattr(
            AppSettings, "WRITER_SESSION_BM25_TOP_K",
            getattr(AppSettings, "WRITER_BM25_TOP_K",
                    getattr(AppSettings, "RETRIEVAL_TOP_K_BM25", 30))
        ))
        req_sim_kb = int(getattr(
            AppSettings, "WRITER_KB_SIM_TOP_K",
            getattr(AppSettings, "WRITER_SIMILARITY_TOP_K",
                    getattr(AppSettings, "RETRIEVAL_TOP_K", 30))
        ))
        req_bm25_kb = int(getattr(
            AppSettings, "WRITER_KB_BM25_TOP_K",
            getattr(AppSettings, "WRITER_BM25_TOP_K",
                    getattr(AppSettings, "RETRIEVAL_TOP_K_BM25", 30))
        ))

        quota_session = int(getattr(AppSettings, "WRITER_SESSION_RETURN_K", 8))
        quota_kb = int(getattr(AppSettings, "WRITER_KB_RETURN_K", 8))
        merged_top_k = int(getattr(AppSettings, "WRITER_MERGED_RETURN_K", 15))
        rrf_k = float(getattr(AppSettings, "WRITER_COMBINE_RRF_K", 60.0))

        # —— 节点视图：会话 & KB（KB 可被“选中集合”过滤） —— #
        session_nodes_all = (sess.nodes if sess else []) or []
        kb_nodes_all = kb.nodes or []

        if use_kb and selected_kb_set:
            kb_nodes_view = [n for n in kb_nodes_all if os.path.basename(self._name_from_meta(n.metadata)) in selected_kb_set]
        else:
            kb_nodes_view = kb_nodes_all if use_kb else []

        logger.info(f"[WriterDBG] nodes: session={len(session_nodes_all)}, kb_total={len(kb_nodes_all)}, kb_view={len(kb_nodes_view)}, kb_selected={len(selected_kb_set)}")

        # 未勾选 use_kb 或筛选后 kb_view 为空时，自动将 KB topk/配额置零
        if not use_kb or len(kb_nodes_view) == 0:
            req_sim_kb = 0
            req_bm25_kb = 0
            quota_kb = 0

        sim_k_sess = self._safe_k(req_sim_sess, len(session_nodes_all), default_min=0)
        bm25_k_sess = self._safe_k(req_bm25_sess, len(session_nodes_all), default_min=0)
        sim_k_kb = self._safe_k(req_sim_kb, len(kb_nodes_view), default_min=0)
        bm25_k_kb = self._safe_k(req_bm25_kb, len(kb_nodes_view), default_min=0)
        logger.info(f"[WriterDBG] effective_topk: sess(sim={sim_k_sess},bm25={bm25_k_sess}) kb(sim={sim_k_kb},bm25={bm25_k_kb})")

        if (sim_k_sess + bm25_k_sess + sim_k_kb + bm25_k_kb) == 0:
            raise ValueError("素材为空：请上传会话文件或（在勾选 KB 后）选择至少一篇知识库文章。")

        session_results: List[NodeWithScore] = []
        kb_results: List[NodeWithScore] = []

        if sess and sess.index and (sim_k_sess > 0 or bm25_k_sess > 0):
            ret_sess = RetrieverFactory.create_hybrid_retriever(
                index=sess.index,
                all_nodes=session_nodes_all,
                similarity_top_k=max(sim_k_sess, 1) if (sim_k_sess + bm25_k_sess) > 0 else 0,
                similarity_top_k_bm25=max(bm25_k_sess, 1) if (sim_k_sess + bm25_k_sess) > 0 else 0
            )
            session_results = ret_sess.retrieve(QueryBundle(query))
        logger.info(f"[WriterDBG] retrieved: session={len(session_results)}")
        self._peek_nodes("retrieved_session", session_results)

        if kb.index and (sim_k_kb > 0 or bm25_k_kb > 0):
            ret_kb = RetrieverFactory.create_hybrid_retriever(
                index=kb.index,
                all_nodes=kb_nodes_view,  # ← BM25 侧仅基于“视图”（被选中文件）
                similarity_top_k=max(sim_k_kb, 1) if (sim_k_kb + bm25_k_kb) > 0 else 0,
                similarity_top_k_bm25=max(bm25_k_kb, 1) if (sim_k_kb + bm25_k_kb) > 0 else 0
            )
            kb_results = ret_kb.retrieve(QueryBundle(query))
            # 二次过滤：若向量侧从全库召回到了不在视图内的节点，则过滤掉
            if selected_kb_set:
                kb_results = [nw for nw in kb_results
                              if os.path.basename(self._name_from_meta((getattr(nw, "node", nw).metadata))) in selected_kb_set]
        logger.info(f"[WriterDBG] retrieved: kb={len(kb_results)} (after filter)")
        self._peek_nodes("retrieved_kb", kb_results)

        if quota_session <= 0 and quota_kb <= 0:
            raise ValueError("来源配额为0：请上调 WRITER_SESSION_RETURN_K/WRITER_KB_RETURN_K。")
        session_slice = session_results[: max(0, quota_session)] if quota_session > 0 else []
        kb_slice = kb_results[: max(0, quota_kb)] if quota_kb > 0 else []
        logger.info(f"[WriterDBG] quota_slice: session={len(session_slice)}/{quota_session}, kb={len(kb_slice)}/{quota_kb}")
        self._peek_nodes("quota_session", session_slice)
        self._peek_nodes("quota_kb", kb_slice)

        merged = RetrieverFactory.rrf_merge(session_slice, kb_slice, k=rrf_k) if hasattr(RetrieverFactory, "rrf_merge") else self._rrf_merge(session_slice, kb_slice, k=rrf_k)
        logger.info(f"[WriterDBG] merged_len={len(merged)} (before final cut)")
        self._peek_nodes("merged", merged)

        if not merged:
            raise ValueError("检索无结果：请检查检索参数或被选 KB 文档是否过少。")

        merged_k = self._safe_k(merged_top_k, len(merged), default_min=1)
        final = merged[:merged_k]
        logger.info(f"[WriterDBG] final_len={len(final)} / merged_top_k={merged_k}")
        self._peek_nodes("final", final)
        return final

    @staticmethod
    def _rrf_merge(list_a: List[NodeWithScore], list_b: List[NodeWithScore], k: float = 60.0) -> List[NodeWithScore]:
        if not list_a:
            return list_b or []
        if not list_b:
            return list_a or []
        rank = {}
        for i, n in enumerate(list_a, start=1):
            nid = n.node.node_id
            rank.setdefault(nid, {"node": n.node, "score": 0.0})
            rank[nid]["score"] += 1.0 / (k + i)
        for i, n in enumerate(list_b, start=1):
            nid = n.node.node_id
            rank.setdefault(nid, {"node": n.node, "score": 0.0})
            rank[nid]["score"] += 1.0 / (k + i)
        merged = sorted(rank.values(), key=lambda x: x["score"], reverse=True)
        return [NodeWithScore(node=it["node"], score=it["score"]) for it in merged]


# 单例
writer_service = WriterService()
